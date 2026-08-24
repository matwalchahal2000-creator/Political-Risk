"""
Pull raw text out of an uploaded annual report / MD&A section or an
earnings-call transcript, regardless of whether it arrives as PDF, DOCX,
or plain text.
"""

import io
from typing import Union

import pdfplumber
from docx import Document


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    file_bytes: raw bytes of the uploaded file (e.g. from Streamlit's
        UploadedFile.getvalue()).
    filename: original filename, used only to pick an extraction path.
    """
    name = filename.lower()

    if name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if name.endswith(".docx"):
        return _extract_docx(file_bytes)
    if name.endswith((".txt", ".md")):
        return _extract_txt(file_bytes)

    # Fall back: try to decode as text; better an attempt than a hard failure.
    try:
        return _extract_txt(file_bytes)
    except Exception as exc:  # pragma: no cover - defensive fallback
        raise ValueError(
            f"Could not read '{filename}': unsupported format and not plain text ({exc})"
        )


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paras = [p.text for p in doc.paragraphs]
    # Tables (common in annual reports) can carry MD&A-adjacent commentary too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text)
    return "\n".join(paras)


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode file as text under utf-8/utf-16/latin-1")
